
# sample of api call
#{"Title":"The Vow","Year":"2012","Rated":"PG-13","Released":"10 Feb 2012","Runtime":"104 min","Genre":"Drama, Romance","Director":"Michael Sucsy","Writer":"Jason Katims, Abby Kohn (screenplay), Stuart Sender (story), Marc Silverstein (screenplay)","Actors":"Rachel McAdams, Channing Tatum, Jessica Lange, Sam Neill","Plot":"A car accident puts Paige in a coma, and when she wakes up with severe memory loss, her husband Leo works to win her heart again.","Language":"English","Country":"USA, France, Australia, UK, Germany, Canada","Awards":"3 wins & 16 nominations.","Poster":"http://ia.media-imdb.com/images/M/MV5BMjE1OTU5MjU0N15BMl5BanBnXkFtZTcwMzI3OTU5Ng@@._V1_SX300.jpg","Metascore":"43","imdbRating":"6.8","imdbVotes":"142,002","imdbID":"tt1606389","Type":"movie","tomatoMeter":"29","tomatoImage":"rotten","tomatoRating":"4.9","tomatoReviews":"126","tomatoFresh":"37","tomatoRotten":"89","tomatoConsensus":"Channing Tatum and Rachel McAdams do their best with what they're given, but The Vow is too shallow and familiar to satisfy the discriminating date-night filmgoer.","tomatoUserMeter":"63","tomatoUserRating":"3.6","tomatoUserReviews":"127536","tomatoURL":"http://www.rottentomatoes.com/m/the_vow_2012/","DVD":"08 May 2012","BoxOffice":"$125.0M","Production":"Sony Pictures","Website":"http://www.thevow-movie.com/","Response":"True"}

def getMovieFeatures(mvName):
    movieApiUrl = "http://www.omdbapi.com/?t="
    apiCall = movieApiUrl + mvName + "&tomatoes=true"
    res = requests.get(apiCall)
    if res.status_code != requests.codes.ok:
        print "Couldn't find info. about " + mvName
        return False
    jsonHndle = res.json()
    print mvName
    trainRecord = []
    for param in InterestingParams:
        print param + " = " + jsonHndle[param]
#        outPara.add_run( " -- " + param + " = " + jsonHndle[param])
        paramVal = jsonHndle[param]
        try:
            featureValue = float(paramVal.replace(',',''))
        except:
            featureValue = 0
        trainRecord.append(featureValue)
    return trainRecord

def readAndGetFeaturesFromDocument(filePath, label):
    inDoc = docx.Document(filePath)
#    outDoc = docx.Document()
    trainFeatures = []
    trainLabels = []    
    for para in inDoc.paragraphs:
        if para.text == "":
            continue
        mvName = para.text
#        outPara = outDoc.add_paragraph(mvName)
        trainRecord = getMovieFeatures(mvName)
        if trainRecord == False:
            continue
        trainFeatures.append(trainRecord)
        trainLabels.append(label)
    print trainFeatures
#    outDoc.save('D:\MoviesRated.docx')
    return trainFeatures,trainLabels
    


import docx
import requests
from sklearn import tree

InterestingParams = {'imdbRating','imdbVotes', 'tomatoRating', 'tomatoReviews', 'tomatoUserMeter' }
trainFeatures,trainLabels = readAndGetFeaturesFromDocument('D:\OutOfWork\Git\MoviesRecommender\Movies.docx', 1)
trainFeatures,trainLabels = readAndGetFeaturesFromDocument('D:\OutOfWork\Git\MoviesRecommender\MoviesRated.docx', 1)

clf = tree.DecisionTreeClassifier()
clf = clf.fit(trainFeatures, trainLabels)
print clf.predict(getMovieFeatures('Beasts of No Nation'))

